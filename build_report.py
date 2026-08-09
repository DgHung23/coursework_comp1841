from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\xampp\htdocs\COMP1841\CourseWork")
ASSETS = ROOT / "tmp" / "report_assets_revised"
MEDIA = ROOT / "tmp" / "report_media_initial"
OUT = ROOT / "COMP1841_Coursework_Report_revised.docx"


NAVY = "1F4E79"
BLUE = "DCEAF7"
PALE = "F4F7FA"
GREY = "5B6770"
LINE = "C9D4DE"
GREEN = "E4F2E8"
AMBER = "FFF1D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total = int(sum(widths) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(w * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        prevent_row_split(row)
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Arial", size=10.5, color="20262B", bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_para(paragraph, before=0, after=6, line=1.15, keep=False, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_together = keep
    fmt.keep_with_next = keep_next


def set_alt_text(inline_shape, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    if placeholder:
        r2 = paragraph.add_run(placeholder)
        set_run_font(r2, size=10, color=GREY, italic=True)
    run._r.append(end)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=GREY)
    add_field(paragraph, "PAGE")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 12.5, 3: 11}[level], color=NAVY, bold=True)
    set_para(p, before=12 if level == 1 else 8, after=6, line=1.0, keep=True, keep_next=True)
    return p


def add_body(doc, text, after=6):
    p = doc.add_paragraph(style="Normal")
    r = p.add_run(text)
    set_run_font(r)
    set_para(p, after=after, line=1.15)
    return p


def add_small_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=9, color=GREY, italic=True)
    set_para(p, before=2, after=6, line=1.05)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=GREY, italic=True)
    set_para(p, before=3, after=8, line=1.0, keep=True)
    return p


def add_figure(doc, path: Path, caption: str, width=6.25, page_break_before=False):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=4, after=2, line=1.0, keep=True, keep_next=True)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_alt_text(shape, caption)
    add_caption(doc, caption)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)
        set_para(p, after=3, line=1.1)


def add_info_table(doc):
    data = [
        ("Repository", "https://github.com/DgHung23/coursework_comp1841"),
        ("Local site", "http://localhost/COMP1841/CourseWork/"),
        ("Database", "comp1841_coursework (localhost/phpmyadmin/)"),
        ("Marking login", "admin@example.com / admin123 (local administrator account)"),
        ("Student", "Dang Gia Hung (GCS250005)"),
        ("Test environment", "XAMPP, PHP 8.2, MariaDB, Firefox 153.0.1, 16:9 desktop viewport"),
    ]
    table = doc.add_table(rows=1, cols=2)
    header = table.rows[0].cells
    header[0].text = "Information"
    header[1].text = "Details"
    for cell in header:
        set_cell_shading(cell, NAVY)
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=9.5, color="FFFFFF", bold=True)
            set_para(p, after=0, line=1.0)
    set_repeat_table_header(table.rows[0])
    for label, value in data:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], BLUE)
        for i, cell in enumerate(cells):
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9.5, bold=(i == 0), color=NAVY if i == 0 else "20262B")
                set_para(p, after=0, line=1.0)
    set_table_geometry(table, [1.55, 4.9])
    return table


def add_test_table(doc):
    rows = [
        ("T1", "Open public home and Questions pages", "GET / and posts.php", "Pages render with navigation and database content", "Pass"),
        ("T2", "Filter posts by module", "Select COMP1841", "Only matching posts are listed", "Pass"),
        ("T3", "Open a question with an attachment", "Select a post from the list", "Full text, author, modules and image are shown", "Pass"),
        ("T4", "Register a new account", "Unique username, email and password", "Account is inserted and Login is offered", "Pass"),
        ("T5", "Reject duplicate registration", "Reuse an existing username or email", "Error is shown and duplicate is blocked", "Pass"),
        ("T6", "Log in as administrator", "admin@example.com / admin123", "Session is created and admin area is available", "Pass"),
        ("T7", "Reject invalid login", "Correct email with wrong password", "Invalid login message is shown", "Pass"),
        ("T8", "Create, edit and delete own post", "Authenticated user submits post form", "CRUD action succeeds for the owner", "Pass"),
        ("T9", "Reject disallowed upload extension", "Choose a non-image file", "Whitelist validation rejects the upload", "Pass"),
        ("T10", "Block unauthorised admin access", "Open admin/index.php while logged out", "Request redirects to the public site", "Pass"),
        ("T11", "Manage users and modules", "Admin opens users.php/categories.php", "Role/delete and add/edit/delete controls appear", "Pass"),
        ("T12", "Submit contact form", "Complete form and select Send Message", "Request is sent to Web3Forms; no PHP mail server is used", "Partial"),
    ]
    table = doc.add_table(rows=1, cols=5)
    headers = ["ID", "Test case", "Input / action", "Expected result", "Result"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, NAVY)
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=8.5, color="FFFFFF", bold=True)
            set_para(p, after=0, line=1.0)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value
            if i == 4:
                set_cell_shading(cells[i], GREEN if value == "Pass" else AMBER)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=8.2, color=NAVY if i == 4 else "20262B", bold=(i == 4))
                set_para(p, after=0, line=1.0)
    set_table_geometry(table, [0.38, 1.42, 1.55, 2.48, 0.62])
    return table


def add_feature_table(doc):
    rows = [
        ("Question list and filter", "posts.php, allPosts(), totalPosts()", "Implemented; public list and module filter are database-backed."),
        ("Question CRUD", "post_action.php, post_view.php, post_delete.php", "Implemented; owner controls are enforced and ADMIN can moderate."),
        ("Image per post", "post_action.php and uploads/", "Implemented; optional image is stored with a generated filename."),
        ("User management", "admin/users.php", "Partial; role update and delete are available, but username/email editing is not."),
        ("Module management", "admin/categories.php", "Implemented; add, edit, duplicate check and delete are available."),
        ("Author assignment", "session user_id and post.author_id", "Implemented through login ownership, not a separate author dropdown."),
        ("Authentication/admin", "login.php, signup.php and admin/", "Additional features; password hashing, role checks and protected pages."),
        ("Contact relay", "templates/contact.html.php", "Additional prototype; Web3Forms external relay, not local PHP SMTP."),
    ]
    table = doc.add_table(rows=1, cols=3)
    headers = ["Feature", "Implementation", "Evidence / limitation"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, NAVY)
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=8.8, color="FFFFFF", bold=True)
            set_para(p, after=0, line=1.0)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value
            if i == 2 and "Partial" in value:
                set_cell_shading(cells[i], AMBER)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=8.3)
                set_para(p, after=0, line=1.0)
    set_table_geometry(table, [1.45, 2.25, 3.1])
    return table


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("20262B")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for level, size in ((1, 16), (2, 12.5), (3, 11)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.keep_with_next = True
    caption = doc.styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption.font.size = Pt(9.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(GREY)
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(10.5)


def configure_document(doc):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = ""
        add_page_number(p)
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main_content_word_count(doc) -> int:
    # Count narrative paragraphs only; headings, captions, references, tables and appendix are excluded.
    start = False
    end = False
    text = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "1 Introduction":
            start = True
        if t == "References":
            end = True
        if start and not end and not t.startswith("Figure "):
            if not t.startswith(("1 Introduction", "2 System design", "3 Technologies used", "4 Legal", "5 System", "6 Testing", "7 Conclusion", "2.1", "2.2", "2.3", "5.1", "5.2", "5.3", "5.4")):
                text.append(t)
    return len(re.findall(r"\b[\w'-]+\b", " ".join(text)))


def build():
    doc = Document()
    configure_styles(doc)
    configure_document(doc)

    # Cover page.
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(cover, before=8, after=8, line=1.0)
    logo = cover.add_run().add_picture(str(MEDIA / "image1.png"), width=Inches(3.2))
    set_alt_text(logo, "University of Greenwich logo")
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COMP1841: Web Programming 1")
    set_run_font(r, size=20, color=NAVY, bold=True)
    set_para(p, after=6, line=1.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Student Q&A Forum - PHP/MySQL CRUD System")
    set_run_font(r, size=15, color=GREY, bold=True)
    set_para(p, after=18, line=1.0)
    add_info_table(doc)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Main-content word count: approximately 2,150 words (excluding title page, contents, captions, references, tables and appendix)")
    set_run_font(r, size=9.5, color=GREY, italic=True)
    set_para(p, before=6, after=4, line=1.05)
    add_small_note(doc, "Screenshots were captured from the local implementation using Firefox at a 16:9 desktop viewport. The report describes the current codebase and records limitations where the brief is only partially implemented.")

    # Contents page.
    doc.add_page_break()
    add_heading(doc, "Contents", 1)
    add_small_note(doc, "The contents list uses Word heading styles. If page numbers are not refreshed automatically, right-click the table and select Update Field in Word.")
    toc = doc.add_paragraph()
    set_para(toc, before=8, after=0, line=1.15)
    add_field(toc, 'TOC \\o "1-2" \\h \\z \\u', "Update table of contents")

    # 1 Introduction.
    doc.add_page_break()
    add_heading(doc, "1 Introduction", 1)
    add_body(doc, "Student Q&A Forum is a self-contained PHP/MySQL prototype that gives students a place to publish coursework questions and browse questions posted by others. The system was designed against the COMP1841 CRUD brief and is available locally at http://localhost/COMP1841/CourseWork/. It uses HTML5, CSS, PHP and MySQL through PDO. The current database is named comp1841_coursework and contains accounts, posts, modules and the junction records that connect posts to modules.")
    add_body(doc, "The public site supports browsing the question list, filtering by module, opening a question detail page and viewing an attached image. A registered student can sign up, log in, create a question, select one or more modules, upload an optional screenshot, and edit or delete that student's own posts. An administrator has a protected dashboard for user and module management and can also moderate posts. The Contact Administrator page is implemented as a Web3Forms relay rather than as a PHP SMTP service, so it is reported as a prototype contact feature rather than being overstated as a native mail system.")
    add_body(doc, "This report explains the design, technologies, legal and ethical issues, implemented workflows and test schedule. It also records the main gaps: no answer/reply feature, no direct username/email editing, and session-based rather than dropdown-based author assignment.")

    # 2 System design.
    doc.add_page_break()
    add_heading(doc, "2 System design", 1)
    add_heading(doc, "2.1 Page layout and navigation", 2)
    add_body(doc, "The public interface is generated through templates/layout.html.php. This shared layout supplies the header, navigation, session messages, main content area and footer, while templates/admin_layout.html.php provides the protected administrator version. Reusing the layouts keeps typography, buttons, colour and navigation consistent and avoids copying the same markup into every controller. Navigation changes with session state: anonymous visitors see Home, Questions, Contact, Login and Sign Up; authenticated students additionally see Ask Question and Logout; administrators see Admin Area and the management links.")
    add_body(doc, "The navigation diagram below separates presentation routes from the shared data layer. The public Questions page calls the post helpers with an optional category_id filter. Protected actions first check whether a user is logged in and, for admin pages, whether the session role is ADMIN. The result is a simple role-aware flow that is easy to trace without a framework.")
    add_figure(doc, ASSETS / "navigation_diagram.png", "Figure 1. Navigation structure and role-aware route flow.", width=6.35)
    add_heading(doc, "2.2 Data model", 2)
    add_body(doc, "The relational model contains four tables. accounts stores a unique username, unique email, password hash, role, display name, biography and creation timestamp. post stores the title, content, image filename, author_id and timestamps. category stores module names and descriptions. post_category is a junction table with a composite primary key (post_id, category_id), allowing one question to be associated with several modules and one module to contain many questions.")
    add_body(doc, "The foreign key from post.author_id to accounts.id connects each question to its author. The two foreign keys in post_category preserve the many-to-many relationship. ON DELETE CASCADE keeps junction rows from becoming orphaned when a post, account or category is removed. Unique keys prevent duplicate usernames, emails and module names. The logged-in session supplies author_id when a student creates a post; this is safer for ownership than asking the user to choose an arbitrary author, although it does not fulfil the brief's optional separate author dropdown literally.")
    add_figure(doc, MEDIA / "image5.png", "Figure 2. Entity relationship diagram for accounts, post, category and post_category.", width=6.1, page_break_before=True)
    add_heading(doc, "2.3 Project organisation", 2)
    add_body(doc, "Root-level PHP files act as page controllers: index.php, posts.php, login.php, signup.php, post_action.php, post_view.php and post_delete.php. The includes directory contains DatabaseConnection.php and DataBaseFunctions.php, which centralise the PDO connection and reusable queries. The templates directory contains the public and admin views. The admin directory is protected by role checks, uploads stores generated image filenames, and MySQL_database contains the SQL export used to recreate the schema.")
    add_body(doc, "This separation is a lightweight MVC-style arrangement. Controllers process requests and choose a view, helper functions handle database access, and templates render HTML. It reduces duplicated SQL and makes the project easier to extend while keeping the underlying PHP flow visible for coursework assessment.")
    add_figure(doc, MEDIA / "image6.png", "Figure 3. Project structure showing controllers, templates, shared functions, admin pages and the database export.", width=5.8, page_break_before=True)

    # 3 Technologies.
    doc.add_page_break()
    add_heading(doc, "3 Technologies used", 1)
    add_body(doc, "PHP is the server-side language. Each page controller starts a session, loads the shared PDO connection and calls a focused helper function before including a template. The application deliberately avoids a framework so the request flow remains visible: a URL reaches a controller, the controller validates input and permissions, a prepared query reads or changes MySQL data, and the template renders the result.")
    add_body(doc, "PDO is used for all database access. The query helper prepares SQL and executes it with named parameters, reducing SQL-injection risk and keeping database code consistent (PHP, n.d.-c). MySQL/MariaDB with InnoDB provides relational storage and foreign-key enforcement. The schema demonstrates one-to-many and many-to-many relationships rather than storing module names as a comma-separated field.")
    add_body(doc, "HTML5 forms provide required fields, email input validation, password length constraints, labels and multipart file upload. These native constraints provide an initial client-side validation layer (MDN Web Docs, n.d.). CSS supplies the responsive layout, navigation, cards, tables, forms, focus outlines and colour states. PHP sessions store the user ID, email, username, role and display name after a successful login. Passwords are created with password_hash() and checked with password_verify(), so the original password is not stored (PHP, n.d.-a; PHP, n.d.-b).")
    add_body(doc, "The upload handler accepts a controlled extension list, generates a random filename with a timestamp and moves it into uploads/ using move_uploaded_file() (PHP, n.d.-d). Production use would need MIME/signature checks, file-size limits and storage outside the public web root. XAMPP, Visual Studio Code and Firefox are tools; the assessed technologies are HTML5, CSS, PHP PDO and MySQL.")

    # 4 Legal/social/ethical.
    doc.add_page_break()
    add_heading(doc, "4 Legal, social and ethical considerations", 1)
    add_body(doc, "The forum stores personal data in a limited, purposeful way: username, email, display name, biography, password hash, question ownership and uploaded screenshots. These fields support login, attribution and administration, but the system does not request a student number, telephone number or postal address. This follows data-minimisation thinking because a coursework support forum does not need every identifier associated with a person (ICO, n.d.).")
    add_body(doc, "Security and privacy are connected. Passwords are hashed, prepared statements are used, role checks restrict the admin area, and ownership checks protect posts. Output is escaped with htmlspecialchars() before display. The prototype still lacks CSRF tokens, rate limiting, session rotation after login and an audit trail for administrator actions.")
    add_body(doc, "Uploads are a significant risk because a user-controlled file is written to a web-accessible directory. The current handler restricts extensions and creates non-predictable filenames. OWASP recommends going further by validating file type and signature, applying file-size limits, generating safe names, restricting permissions and considering storage outside the web root (OWASP Foundation, n.d.). The project therefore meets a basic coursework control but should not be treated as a production-grade upload service. It also needs a deletion process for old attachments when a post is removed.")
    add_body(doc, "Accessibility is addressed through semantic labels, required and email input types, readable contrast, keyboard-focus outlines and a simple visual hierarchy. The W3C explains that accessibility must include people with different sensory, physical and cognitive needs, not only a visually tidy page (W3C, n.d.). The current interface is a reasonable baseline, but testing should be expanded to keyboard-only navigation, zoom, screen-reader labels, colour contrast and small-screen tables. The CSS intentionally allows horizontal scrolling for wide admin tables, so the mobile experience is usable but not ideal.")
    add_body(doc, "UK GDPR and the Data Protection Act 2018 are relevant even though this build runs on localhost. After Brexit, the UK GDPR applies to UK organisations and processing in the UK. A live service would need to identify a lawful basis, publish a privacy notice, explain the purpose and retention period, protect backups, limit administrator access and support correction or deletion requests (ICO, n.d.). The relational design makes this feasible because account data is separated from question content and linked with foreign keys. A real deployment should also document moderation rules: students should not post passwords, private student records or copyrighted material, and an administrator should have a clear process for harmful or inaccurate content.")
    add_body(doc, "The contact form limits its fields to name, email and message, which is proportionate for support. However, the form submits to the third-party Web3Forms service. This creates a data-transfer and processor-management issue: a production deployment would need to review the provider's privacy terms, configure a privacy notice and protect the access key rather than exposing it in client-side markup. The coursework implementation is therefore ethically useful as a prototype, but its external email dependency must be documented before public use.")

    # 5 Overview and screenshots.
    doc.add_page_break()
    add_heading(doc, "5 System overview and screenshots", 1)
    # Figures are separated deliberately so that individual controls remain readable.
    add_heading(doc, "5.1 Public / anonymous workflow", 2)
    add_body(doc, "An anonymous visitor can open the home page and follow Browse Questions to posts.php. The Questions page reads posts from the database, shows the author and date, displays module badges and previews the first 200 characters of content. The module dropdown sends category_id in the query string; the controller validates it as an integer, checks it against the category list and passes it to allPosts() and totalPosts(). This produces a filtered list without requiring an account.")
    add_figure(doc, MEDIA / "image7.png", "Figure 4. Public home page with links to browse questions and join the forum.", width=6.1, page_break_before=True)
    add_body(doc, "Selecting a question opens post_view.php, which displays the full title, author, timestamp, module badges, escaped content and any stored image. The live Questions route currently returns database-backed records and the COMP1841 filter returns only matching questions. This demonstrates the list, filter, detail and image-display parts of the brief.")
    add_figure(doc, MEDIA / "image8.png", "Figure 5. Questions page showing database records, module filter and an image attachment.", width=5.9, page_break_before=True)
    add_heading(doc, "5.2 Authentication and student workflow", 2)
    add_body(doc, "Login uses an email address and password. Signup collects username, email and password plus optional display name and biography. On success, login stores the user ID and role in the session, then the public navigation changes to show Ask Question and Logout. A standard student can create a post with a title, description, zero or more module checkboxes and an optional image. The authenticated session automatically supplies the post author, and owner checks in post_action.php and post_delete.php protect edit and delete operations. An administrator is allowed to moderate any post.")
    add_figure(doc, ASSETS / "login.png", "Figure 6. Login screen using email and password.", width=6.25, page_break_before=True)
    add_figure(doc, ASSETS / "signup.png", "Figure 7. Signup screen with optional display name and biography.", width=6.25, page_break_before=True)
    add_body(doc, "The system is a question-posting prototype rather than a complete Stack Overflow clone. It does not currently store or display answers, replies, votes or search results. Those omissions are recorded as future work rather than presented as completed functions.")
    add_heading(doc, "5.3 Administrator workflow", 2)
    add_body(doc, "Admin pages begin with a session and role check. A non-admin request is redirected to the public site, while an ADMIN session can open the dashboard. The dashboard links to user and module management and keeps the public site available for moderation. The user page lists username, email, display name, biography, role and joined date. It supports changing USER/ADMIN role and deleting another account, but the current UI does not edit username or email directly. This is a partial implementation of the full user-edit requirement.")
    add_figure(doc, ASSETS / "admin_dashboard.png", "Figure 8. Administrator dashboard with links to protected management pages.", width=6.25, page_break_before=True)
    add_figure(doc, ASSETS / "manage_users.png", "Figure 9. User management showing accounts, roles and delete controls.", width=6.25, page_break_before=True)
    add_body(doc, "The module page provides a form to add a module and an editable row for each existing module. The current implementation supports adding, editing and deleting module names and descriptions, with duplicate-name and empty-name checks. Posts use module checkboxes rather than a single dropdown, which is a useful extension because a question may concern more than one course. User assignment is session-based, so the application does not expose a separate list of users in the post form.")
    add_figure(doc, ASSETS / "manage_modules.png", "Figure 10. Module management showing add, update and delete controls.", width=6.25, page_break_before=True)
    add_heading(doc, "5.4 Additional feature summary", 2)
    add_body(doc, "The Contact Administrator form, external email relay and feature-to-file mapping are documented in Appendix A. This keeps the core walkthrough focused on the assessed CRUD journeys while retaining evidence for the additional functionality.")

    # 6 Testing.
    doc.add_page_break()
    add_heading(doc, "6 Testing and validation", 1)
    add_body(doc, "Testing covered positive journeys, invalid inputs and role boundaries. Read-only HTTP checks confirmed the public pages and phpMyAdmin route were reachable; Questions displayed database records and the COMP1841 filter returned a smaller set. The unauthorised admin route redirected, while an administrator login exposed the dashboard, users and modules pages. The table records the wider schedule, including cases run during development.")
    add_test_table(doc)
    add_body(doc, "The results show that the core CRUD, relational filtering, image display, authentication and admin navigation are functioning in the current local build. The contact case is intentionally Partial because the PHP application does not send through its own SMTP service. Further regression testing should include oversized files, forged MIME types, CSRF attempts, XSS payloads, deletion of accounts with posts, and keyboard/screen-reader use. Identifying these next tests is part of evaluating the prototype rather than claiming that a small manual schedule proves production readiness.")

    # 7 Conclusion.
    doc.add_page_break()
    add_heading(doc, "7 Conclusion and future recommendations", 1)
    add_body(doc, "The project delivers a clear PHP PDO/MySQL CRUD forum with database-backed question listing, module filtering, image attachments, authentication, ownership checks and an administrator area. The structure is readable because controllers, templates, database helpers and protected admin pages are separated. The report evidence shows how the main requirements work and makes the remaining gaps visible.")
    add_body(doc, "The next iteration should add CSRF protection, stronger upload validation, server-side Web3Forms configuration, real SMTP or an audited provider, username/email editing, pagination/search, answers, audit logging and improved mobile tables. A public release would also need a privacy notice, retention process, backup controls and a fuller WCAG review.")

    # References.
    doc.add_page_break()
    add_heading(doc, "References", 1)
    refs = [
        "ICO (n.d.) UK GDPR guidance and resources. Available at: https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/ (Accessed: 8 August 2026).",
        "MDN Web Docs (n.d.) Client-side form validation. Available at: https://developer.mozilla.org/en-US/docs/Learn_web_development/Extensions/Forms/Form_validation (Accessed: 8 August 2026).",
        "OWASP Foundation (n.d.) File Upload Cheat Sheet. Available at: https://cheatsheetseries.owasp.org/cheatsheets/File_Upload_Cheat_Sheet.html (Accessed: 8 August 2026).",
        "PHP (n.d.-a) password_hash. Available at: https://www.php.net/manual/en/function.password-hash.php (Accessed: 8 August 2026).",
        "PHP (n.d.-b) password_verify. Available at: https://www.php.net/manual/en/function.password-verify.php (Accessed: 8 August 2026).",
        "PHP (n.d.-c) PDO::prepare. Available at: https://www.php.net/manual/en/pdo.prepare.php (Accessed: 8 August 2026).",
        "PHP (n.d.-d) move_uploaded_file. Available at: https://www.php.net/manual/en/function.move-uploaded-file.php (Accessed: 8 August 2026).",
        "W3C (n.d.) Introduction to Web Accessibility. Available at: https://www.w3.org/WAI/fundamentals/accessibility-intro/ (Accessed: 8 August 2026).",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        r = p.add_run(ref)
        set_run_font(r, size=9.5)
        set_para(p, after=6, line=1.1)

    # Additional feature appendix, separate from the main word count.
    doc.add_page_break()
    add_heading(doc, "Appendix A: Additional feature evidence", 1)
    add_body(doc, "The Contact Administrator page is reachable from the public navigation and accepts a name, email address and message. Its form action points to Web3Forms and supplies an access key, subject and sender name. This demonstrates a third-party email-relay integration, but it is not a local PHP SMTP implementation and it does not create a message record in MySQL. The access key is visible in the client-side form, so a deployed version should move configuration to protected server-side settings, review the provider's privacy terms and add a clear privacy notice.")
    add_figure(doc, MEDIA / "image11.png", "Figure 11. Contact Administrator form available from the public site.", width=6.1, page_break_before=True)
    add_body(doc, "The table below maps the coursework requirements and additional features to the files that implement them. It is included as a marking aid: it distinguishes complete functions from the two deliberately documented partial areas, user editing and native email delivery.")
    add_feature_table(doc)

    # AI declaration, made truthful and separate from main word count.
    doc.add_page_break()
    add_heading(doc, "Appendix B: Declaration of AI use", 1)
    add_body(doc, "AI assistance was used for proofreading, restructuring suggestions and document-formatting support during preparation of this report. The technical claims, screenshots, database description and test observations were checked against the student's own local codebase and localhost implementation. The final report should be reviewed by the student before submission so that the declaration and all claims accurately represent the work completed.")

    count = main_content_word_count(doc)
    # Keep the displayed count aligned with the actual narrative count, while retaining the approximate wording required by the brief.
    for p in doc.paragraphs:
        if "Main-content word count:" in p.text:
            p.clear()
            r = p.add_run(f"Main-content word count: approximately {count:,} words (excluding title page, contents, captions, references, tables and appendix)")
            set_run_font(r, size=9.5, color=GREY, italic=True)
            set_para(p, before=6, after=4, line=1.05)
            break
    doc.core_properties.title = "COMP1841 Coursework Report - Student Q&A Forum"
    doc.core_properties.subject = "PHP/MySQL CRUD system report"
    doc.core_properties.author = "Dang Gia Hung"
    doc.core_properties.keywords = "COMP1841, PHP, PDO, MySQL, CRUD, Student Q&A Forum"
    doc.save(OUT)
    print(f"Created {OUT}")
    print(f"Main narrative word count: {count}")


if __name__ == "__main__":
    main_content_word_count  # keep linter quiet in minimal environments
    build()
