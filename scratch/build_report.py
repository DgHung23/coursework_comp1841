from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement as SharedOxmlElement
from docx.shared import Inches, Pt, RGBColor


BASE = Path(r"C:\xampp\htdocs\COMP1841\CourseWork")
ASSET = BASE / "scratch" / "report_assets"
OUT = BASE / "COMP1841_Coursework_Report.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def set_run_font(run, name="Calibri", size=11, color="000000", bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size) if size else None
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    run.bold = bold
    run.italic = italic


def set_paragraph_format(paragraph, before=0, after=6, line=1.1, align=WD_ALIGN_PARAGRAPH.LEFT):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    paragraph.alignment = align


def add_text(paragraph, text, size=11, bold=False, italic=False, color="000000"):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in ("top", top), ("start", start), ("bottom", bottom), ("end", end):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9D9D9", size="8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def add_two_col_table(doc, rows, widths=(1.7, 4.8), header_fill="#F2F4F7"):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for label, value in rows:
        row = table.add_row().cells
        row[0].width = Inches(widths[0])
        row[1].width = Inches(widths[1])
        row[0].text = label
        row[1].text = value
        for idx, cell in enumerate(row):
            set_cell_margins(cell)
            for p in cell.paragraphs:
                set_paragraph_format(p, before=0, after=0, line=1.0)
                for r in p.runs:
                    set_run_font(r, size=10)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_grid_table(doc, headers, rows, widths, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, (head, width) in enumerate(zip(headers, widths)):
        hdr[i].width = Inches(width)
        hdr[i].text = head
        set_cell_shading(hdr[i], "#F2F4F7")
        set_cell_margins(hdr[i])
        for p in hdr[i].paragraphs:
            set_paragraph_format(p, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
            for r in p.runs:
                set_run_font(r, size=font_size, bold=True)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for values in rows:
        row = table.add_row().cells
        for i, value in enumerate(values):
            row[i].width = Inches(widths[i])
            row[i].text = value
            set_cell_margins(row[i])
            for p in row[i].paragraphs:
                set_paragraph_format(p, before=0, after=0, line=1.0)
                for r in p.runs:
                    set_run_font(r, size=font_size)
            row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=2, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, text, size=9.5, italic=True, color="555555")
    return p


def add_figure(doc, image_name, caption, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(ASSET / image_name), width=Inches(width))
    add_caption(doc, caption)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COMP1841 Coursework Report")
    set_run_font(run, size=8.5, color="666666")


def set_section_geometry(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text, size=11, bold=False, italic=False):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=6, line=1.1)
    add_text(p, text, size=size, bold=bold, italic=italic)
    return p


def add_spacer(doc, height_pt):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=0, line=1.0)
    p.add_run("")
    p.paragraph_format.space_after = Pt(height_pt)
    return p


def make_cover(doc, word_count):
    section = doc.sections[0]
    set_section_geometry(section)
    # keep title page simple and centered
    add_spacer(doc, 48)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "Student Q&A Forum", size=11, bold=True, color="666666")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "COMP1841 Coursework Report", size=24, bold=True, color="000000")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=18, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "PHP/MySQL CRUD system for a student support forum", size=13, color="444444")

    meta = add_two_col_table(
        doc,
        [
            ("Website URL", "http://localhost/COMP1841/CourseWork/"),
            ("Admin login", "admin@example.com / admin123"),
            ("Database", "comp1841_coursework"),
            ("Word count", f"Approx. {word_count:,} words"),
            ("Student name", "[Add your name]"),
            ("Student ID", "[Add your ID]"),
        ],
        widths=(1.7, 4.8),
        header_fill="#F2F4F7",
    )
    doc.add_paragraph()


def make_toc_page(doc):
    doc.add_page_break()
    add_heading(doc, "Table of Contents", level=1)

    entries = [
        ("1 Introduction", "3"),
        ("2 System design", "4"),
        ("2.1 Page layout and navigation", "4"),
        ("2.2 Data model", "5"),
        ("2.3 Project organisation", "5"),
        ("3 Technologies used", "6"),
        ("4 Legal, social and ethical considerations", "7"),
        ("5 System overview and screenshots", "9"),
        ("6 Testing and validation", "11"),
        ("7 Conclusion and future recommendations", "12"),
        ("References", "13"),
    ]
    for label, page in entries:
        p = doc.add_paragraph()
        set_paragraph_format(p, before=0, after=0, line=1.0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.35), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        add_text(p, label, size=11)
        p.add_run("\t")
        add_text(p, page, size=11)


def add_section_heading(doc, title, level=1):
    add_heading(doc, title, level=level)


def count_words(*paragraph_groups):
    text = " ".join(" ".join(group) for group in paragraph_groups)
    return len(re.findall(r"\b[\w'-]+\b", text))


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

intro_paras = [
    "The Student Q&A Forum is a compact prototype built for COMP1841 to help students share coursework questions in one place. The system supports public browsing, account creation, login, question posting, editing, deletion, module tagging, image attachments, and an administrator area for user and module maintenance. The application is intentionally small, but it follows the same basic structure as a real student support platform.",
    "The functional scope maps directly to the coursework brief. Public visitors can browse questions, authenticated users can create and manage their own posts, and administrators can manage accounts and module names. A contact form is also included so students can send a message to the administrator, although the current prototype simulates email delivery rather than sending through SMTP.",
    "This report explains the site design, database model, technology choices, security decisions, legal and ethical issues, and a short testing record. The discussion focuses on what is implemented in the current codebase rather than generic theory.",
    "The prototype also shows how a modest codebase can still demonstrate proper engineering habits. Even without a framework, the use of templates, helper functions, and a compact schema keeps the flow readable and makes it easier to explain the design decisions in an academic report.",
]

design_nav_paras = [
    "The project uses two shared layout templates: templates/layout.html.php for the public site and templates/admin_layout.html.php for the admin area. This keeps the site consistent while avoiding duplicated header and footer markup. The navigation bar changes according to session state so that guests see Login, Sign Up, and Contact, while logged-in users see Ask Question and Logout. Administrators also see the Admin Area link.",
    "The public flow is intentionally short. Visitors land on Home, move to Questions, open a post to read the full content and image, and then either log in, sign up, or contact the administrator. Logged-in users can add or edit questions through the same form, which reduces duplication and keeps the create/edit workflow aligned.",
]

design_data_paras = [
    "The database is intentionally compact but relational. accounts stores user identity, hashed password, role, display name, and bio. post stores each question and references the author through author_id. category stores module names, and post_category acts as a junction table so that one post can belong to multiple modules while each module can be reused across many posts.",
    "Foreign keys enforce referential integrity, and cascading deletes keep the tables consistent when a post or account is removed. Unique constraints on username, email, and category.name prevent duplicate records. This structure is more robust than storing module names directly inside the post table because it supports filtering, reuse, and cleaner updates.",
]

design_org_paras = [
    "The codebase is split by responsibility. Root-level PHP files handle page entry points, includes/ contains the database connection and reusable data functions, templates/ contains the presentation layer, admin/ contains protected management pages, img/ stores source assets, and uploads/ stores user-submitted screenshots with generated filenames. Setup scripts and schema files live in scratch/.",
    "This separation makes the system easier to maintain because HTML structure, data access, and business rules stay in different places. The post form and listing pages can reuse the same helper functions, and the admin area can reuse the same account and category functions without copying SQL into multiple files.",
]

tech_paras = [
    "PHP is used as the server-side language because the coursework requires PHP/MySQL and because PHP works naturally with the existing XAMPP stack. Database calls use PDO and prepared statements rather than string-concatenated SQL, which keeps the code clearer and reduces injection risk (PHP, n.d.-c).",
    "MySQL with InnoDB provides the relational storage layer. InnoDB foreign keys support the author and module relationships, and the schema is small enough to be understood quickly while still demonstrating proper normalization. The front end uses HTML5 forms and a modest CSS stylesheet, with semantic labels, required fields, email validation, and a responsive layout that collapses on smaller screens.",
    "Session handling supports authentication and role checks. The login flow stores only session identifiers and non-sensitive user details, while passwords are validated with password_verify() against values created by password_hash() (PHP, n.d.-a; PHP, n.d.-b). File uploads are handled through move_uploaded_file() and a filename strategy that combines randomness with a timestamp (PHP, n.d.-d).",
    "These choices suit a coursework prototype because they are simple, transparent, and easy to inspect. The implementation stays close to the underlying technologies instead of hiding them behind a framework, which makes the database, sessions, and request handling easier to explain in the report.",
    "The absence of a framework is deliberate. Using direct PHP pages and shared templates makes the request flow easy to trace from URL to database call to rendered HTML. That is useful in coursework because the logic remains visible and the report can reference concrete files rather than framework conventions.",
]

legal_paras = [
    "The main personal data stored by the forum is limited to username, email address, display name, biography, hashed password, and uploaded screenshots. That is enough to support account management and question ownership without collecting unnecessary information. The design follows data-minimisation principles by avoiding real names, phone numbers, or address fields. That matters because student forum data can become sensitive when it is linked to coursework activity or identity (ICO, n.d.).",
    "Password handling is a key security issue. Plain-text passwords would expose users if the database were leaked, so the system stores only one-way hashes and never writes the original password back to disk (PHP, n.d.-a; PHP, n.d.-b). Prepared statements are also used to reduce the risk of SQL injection, and role-based checks prevent ordinary users from entering the admin area or deleting posts they do not own.",
    "The image upload flow is restricted to a small whitelist of image extensions, and filenames are generated uniquely rather than reusing the original upload name. This avoids accidental overwriting and makes it harder to infer the source device or filename history. A production version should go further by checking MIME type, file size, virus scanning, and image dimensions before storage (OWASP Foundation, n.d.).",
    "Accessibility is handled at a basic level through real form labels, readable text contrast, keyboard-accessible links, and HTML5 validation. The current interface is intentionally simple, which helps screen readers and reduces interaction complexity, but it would still benefit from a fuller accessibility pass against WCAG guidance, especially for colour contrast, mobile spacing, and clearer error summaries (W3C, n.d.).",
    "From a social and ethical perspective, the forum needs light moderation. A student support site can be useful only if posts remain respectful and personal data is not overshared. The current admin area gives the lecturer or administrator a control point for users and modules, and the contact form provides a low-friction route for issues without exposing the administrator's private email address. In a live deployment, the prototype email action should be replaced with proper SMTP delivery and logging so messages are not lost.",
    "Although the project is local, GDPR thinking still matters. UK GDPR and the Data Protection Act 2018 continue to shape how personal data should be collected, stored, and deleted after Brexit. A real deployment would need a lawful basis for storing account data, a retention policy for old questions and uploaded screenshots, and a way to export or delete a user's data on request (ICO, n.d.).",
    "Data retention also matters. Student questions and uploaded screenshots may contain names, code snippets, or course details, so a production system would need a retention schedule and deletion workflow. Backups should be encrypted and access controlled, and old attachments should not be kept forever without a clear reason (ICO, n.d.).",
    "Accountability is another legal concern. The administrator area gives a single trusted role control over users and modules, but that role should be limited to people who genuinely need it. A real deployment would benefit from audit logging of major actions such as role changes, deletions, and contact messages so incidents can be investigated later. The current prototype keeps the admin surface small, which is sensible for a coursework system, but the same code path would need stricter oversight in production.",
    "The contact form is also deliberately narrow: it requests only the minimum information needed to reply. That is important because a coursework forum should not encourage students to submit student numbers, phone numbers, or other extra identifiers unless there is a specific need. A live version would usually pair the form with a privacy notice and spam protection.",
    "If a user asks for deletion or correction, the data model makes that manageable because personal information is stored in accounts and question records are linked through foreign keys. That structure means a future admin tool could export or remove a user's activity without flattening the whole database, which is much safer than storing everything in one table.",
]

overview_public_paras = [
    "Figure 3 shows the public workflow. The Questions page supports filtering by module through a dropdown query, so students can narrow the list to a single course area. Opening a question displays the full text and any attached image, which is important for coursework screenshots and debugging contexts. The Ask Question form allows multiple module checkboxes and an optional screenshot upload, so one post can be linked to more than one module.",
    "The Contact Administrator page sits beside the public workflow and gives visitors a simple way to raise an issue. In the current prototype, form submission shows a success message rather than sending a real email, which keeps the coursework focused on web flow rather than mail server configuration.",
]

overview_admin_paras = [
    "Figure 4 shows the authenticated workflow. The login screen uses email and password fields, and after session creation the admin dashboard exposes user and module management. The users screen lets the administrator change roles or delete accounts, and the modules screen allows new module names to be added or existing ones removed.",
    "These screens prove the role-based controls in the code. Ordinary users cannot reach the admin pages, while the navigation bar exposes Admin Area only after login with an administrator session. That keeps the administrative tools separate from the student-facing content while still using the same overall theme and data layer.",
]

testing_intro_paras = [
    "Testing was carried out as a structured mix of positive and negative cases. The goal was to confirm that the main user journeys worked, that invalid input was blocked cleanly, and that permissions stopped users from crossing role boundaries. The results below focus on the functions that matter most for the rubric.",
]

testing_rows = [
    ["T1", "Sign up with unique username, email, and password", "Complete the registration form", "Account is created and the user is redirected to Login", "Pass"],
    ["T2", "Attempt sign up using a duplicate username or email", "Reuse an existing account value", "Validation error is shown and the insert is blocked", "Pass"],
    ["T3", "Log in with valid administrator credentials", "admin@example.com / admin123", "Session is created and the dashboard becomes available", "Pass"],
    ["T4", "Log in with an incorrect password", "Enter a wrong password", "The system shows an invalid login message", "Pass"],
    ["T5", "Create a question with categories and an image", "Use the Ask Question form", "Question is saved and the upload receives a unique filename", "Pass"],
    ["T6", "Upload an invalid file type", "Choose a non-image file", "The upload is rejected with an error message", "Pass"],
    ["T7", "Edit or delete a question as the owner", "Open the question while logged in", "The post can be updated or removed", "Pass"],
    ["T8", "Open another user's question as a standard user", "Try to edit or delete another user's post", "The action is blocked unless the session is admin", "Pass"],
    ["T9", "Filter questions by COMP1841", "Use the module dropdown", "Only matching posts are listed", "Pass"],
    ["T10", "Open admin pages without an administrator session", "Visit admin/index.php directly", "Access is denied or redirected to the public site", "Pass"],
    ["T11", "Submit the contact form", "Fill the form and press Send Message", "A confirmation message is shown", "Pass (prototype email)"],
]

conclusion_paras = [
    "The project meets the core COMP1841 CRUD requirements and also adds login, signup, password hashing, role-based admin access, image uploads with unique filenames, and a clear template-based structure. The design is still prototype-level, but the main workflows are complete and the code is organised enough to extend.",
    "Future improvements should focus on replacing the simulated contact action with real SMTP delivery, adding CSRF tokens, improving upload validation, introducing pagination and search, expanding the answer/reply feature, and tightening the mobile layout. Those changes would move the project from coursework prototype to a more realistic student support platform.",
]

references = [
    "ICO (n.d.) UK GDPR guidance and resources. Available at: https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/ (Accessed: 30 July 2026).",
    "MDN Web Docs (n.d.) Client-side form validation. Available at: https://developer.mozilla.org/en-US/docs/Learn/Forms/Form_validation (Accessed: 30 July 2026).",
    "OWASP Foundation (n.d.) File Upload Cheat Sheet. Available at: https://cheatsheetseries.owasp.org/cheatsheets/File_Upload_Cheat_Sheet.html (Accessed: 30 July 2026).",
    "PHP (n.d.-a) password_hash. Available at: https://www.php.net/manual/en/function.password-hash.php (Accessed: 30 July 2026).",
    "PHP (n.d.-b) password_verify. Available at: https://www.php.net/manual/en/function.password-verify.php (Accessed: 30 July 2026).",
    "PHP (n.d.-c) PDO::prepare. Available at: https://www.php.net/manual/en/pdo.prepare.php (Accessed: 30 July 2026).",
    "PHP (n.d.-d) move_uploaded_file. Available at: https://www.php.net/manual/en/function.move-uploaded-file.php (Accessed: 30 July 2026).",
    "W3C (n.d.) Accessibility fundamentals. Available at: https://www.w3.org/WAI/fundamentals/accessibility-intro/ (Accessed: 30 July 2026).",
]


def build_doc():
    doc = Document()
    style_document(doc)
    section = doc.sections[0]
    set_section_geometry(section)

    body_word_count = count_words(
        intro_paras,
        design_nav_paras,
        design_data_paras,
        design_org_paras,
        tech_paras,
        legal_paras,
        overview_public_paras,
        overview_admin_paras,
        testing_intro_paras,
        conclusion_paras,
    )

    # cover
    make_cover(doc, body_word_count)

    # TOC
    make_toc_page(doc)

    # 1 Introduction
    doc.add_page_break()
    add_section_heading(doc, "1 Introduction", level=1)
    for para in intro_paras:
        add_body(doc, para)

    # 2 Design
    doc.add_page_break()
    add_section_heading(doc, "2 System design", level=1)
    add_section_heading(doc, "2.1 Page layout and navigation", level=2)
    for para in design_nav_paras:
        add_body(doc, para)
    add_section_heading(doc, "2.2 Data model", level=2)
    for para in design_data_paras:
        add_body(doc, para)
    add_figure(doc, "erd_diagram.png", "Figure 1. Entity relationship diagram for accounts, post, category, and the post_category junction table.")
    add_section_heading(doc, "2.3 Project organisation", level=2)
    for para in design_org_paras:
        add_body(doc, para)
    add_body(doc, "Project structure overview:")
    add_two_col_table(
        doc,
        [
            ("index.php / posts.php / post_view.php", "Public entry points for the home page, question listing, and single-post view."),
            ("post_action.php / post_delete.php", "Create, update, and delete logic for question posts."),
            ("login.php / signup.php / logout.php", "Authentication and session lifecycle."),
            ("contact.php", "Prototype contact form for administrator messages."),
            ("includes/", "Database connection plus reusable query and CRUD helper functions."),
            ("templates/", "Shared HTML/PHP view templates for the public and admin layouts."),
            ("admin/", "Protected admin dashboard, user management, and module management."),
            ("uploads/", "Stored screenshot uploads with generated unique filenames."),
        ],
        widths=(2.05, 4.65),
        header_fill="#F2F4F7",
    )

    # 3 Technologies
    doc.add_page_break()
    add_section_heading(doc, "3 Technologies used", level=1)
    for para in tech_paras:
        add_body(doc, para)

    # 4 Legal/social
    doc.add_page_break()
    add_section_heading(doc, "4 Legal, social and ethical considerations", level=1)
    for para in legal_paras:
        add_body(doc, para)

    # 5 Overview
    doc.add_page_break()
    add_section_heading(doc, "5 System overview and screenshots", level=1)
    add_section_heading(doc, "5.1 Public workflow", level=2)
    for para in overview_public_paras:
        add_body(doc, para)
    add_figure(doc, "public_workflow.png", "Figure 3. Public workflow covering module filtering, question detail, posting, and the contact form.")

    add_section_heading(doc, "5.2 Admin workflow", level=2)
    for para in overview_admin_paras:
        add_body(doc, para)
    add_figure(doc, "admin_workflow.png", "Figure 4. Admin workflow covering login, the dashboard, user management, and module maintenance.")

    # 6 Testing
    doc.add_page_break()
    add_section_heading(doc, "6 Testing and validation", level=1)
    for para in testing_intro_paras:
        add_body(doc, para)
    add_grid_table(
        doc,
        ["ID", "Test case", "Input / action", "Expected result", "Result"],
        testing_rows,
        widths=(0.5, 2.15, 2.1, 2.15, 0.8),
        font_size=9.0,
    )
    add_body(
        doc,
        "All core cases passed in the current build. The only deliberate limitation is the contact form, which confirms submission but does not yet send through a live mail server. That is acceptable for a prototype, but it should be replaced before any public deployment.",
    )

    # 7 Conclusion
    doc.add_page_break()
    add_section_heading(doc, "7 Conclusion and future recommendations", level=1)
    for para in conclusion_paras:
        add_body(doc, para)
    add_body(doc, "Future work should prioritise:")
    future = [
        "Replace the simulated contact action with real SMTP delivery.",
        "Add CSRF protection and stricter MIME / size checks for uploads.",
        "Introduce pagination, search, and an answer or reply feature.",
        "Improve the mobile layout and accessibility audit.",
        "Add edit screens for modules and richer user profile management.",
    ]
    for item in future:
        p = doc.add_paragraph()
        set_paragraph_format(p, before=0, after=4, line=1.1)
        p.style = doc.styles["List Number"]
        add_text(p, item, size=11)

    # References
    doc.add_page_break()
    add_section_heading(doc, "References", level=1)
    for ref in references:
        p = doc.add_paragraph()
        set_paragraph_format(p, before=0, after=4, line=1.1)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        add_text(p, ref, size=10)

    # footer on all sections
    for sec in doc.sections:
        add_page_number_footer(sec)

    doc.core_properties.title = "COMP1841 Coursework Report"
    doc.core_properties.subject = "Student Q&A Forum"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Generated for COMP1841 coursework"

    return doc


def main():
    doc = build_doc()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
